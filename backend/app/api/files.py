import os
import uuid
import tempfile
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
import aiofiles

# Импорты для прямой конвертации Office файлов в изображения через Python
try:
    from docx import Document
    from PIL import Image, ImageDraw, ImageFont
    import openpyxl
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    OFFICE_CONVERSION_AVAILABLE = True
    print("✅ Библиотеки для прямой конвертации Office файлов доступны")
except ImportError as e:
    OFFICE_CONVERSION_AVAILABLE = False
    print(f"⚠️ Библиотеки для конвертации Office файлов недоступны: {e}")

from ..database import get_db
from ..models.request import Request
from ..models.request_file import RequestFile
from ..models.user import User
from ..schemas.request_file import RequestFileResponse
from ..dependencies import get_current_user

router = APIRouter()

# Папка для хранения файлов
UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

# Папка для кеша изображений
IMAGE_CACHE_DIR = "image_cache"
if not os.path.exists(IMAGE_CACHE_DIR):
    os.makedirs(IMAGE_CACHE_DIR)

def convert_docx_to_image(docx_path: str) -> str:
    """Конвертирует DOCX документ в изображение через python-docx"""
    
    file_filename = os.path.basename(docx_path)
    base_name = os.path.splitext(file_filename)[0]
    
    print(f"📝 Конвертируем DOCX документ в изображение: {docx_path}")
    
    try:
        # Читаем DOCX файл
        doc = Document(docx_path)
        
        # Настройки изображения
        page_width = 800
        page_height = 1100
        margin = 60
        content_width = page_width - 2 * margin
        
        # Создаем изображение
        image = Image.new('RGB', (page_width, page_height), 'white')
        draw = ImageDraw.Draw(image)
        
        # Пытаемся загрузить шрифты
        try:
            font_normal = ImageFont.truetype("arial.ttf", 12)
            font_bold = ImageFont.truetype("arialbd.ttf", 14)
            font_heading = ImageFont.truetype("arialbd.ttf", 16)
        except:
            font_normal = ImageFont.load_default()
            font_bold = ImageFont.load_default()
            font_heading = ImageFont.load_default()
        
        y_offset = margin
        
        # Рисуем параграфы
        for para in doc.paragraphs:
            if para.text.strip():
                # Выбираем шрифт в зависимости от стиля
                font = font_normal
                if para.style.name.startswith('Heading'):
                    font = font_heading
                elif any(run.bold for run in para.runs if run.bold):
                    font = font_bold
                
                # Разбиваем длинный текст на строки
                words = para.text.split()
                lines = []
                current_line = []
                
                for word in words:
                    test_line = ' '.join(current_line + [word])
                    if len(test_line) * 7 <= content_width:  # Примерная ширина символа
                        current_line.append(word)
                    else:
                        if current_line:
                            lines.append(' '.join(current_line))
                            current_line = [word]
                        else:
                            lines.append(word)
                
                if current_line:
                    lines.append(' '.join(current_line))
                
                # Рисуем строки
                for line in lines:
                    if y_offset > page_height - margin:
                        break  # Достигли конца страницы
                    
                    draw.text((margin, y_offset), line, fill='black', font=font)
                    y_offset += 20
                
                y_offset += 10  # Отступ между параграфами
        
        # Обрабатываем таблицы
        for table in doc.tables:
            if y_offset > page_height - margin - 100:
                break  # Не помещается таблица
            
            # Простая отрисовка таблицы
            col_width = content_width // len(table.columns) if table.columns else content_width
            row_height = 25
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    x = margin + col_idx * col_width
                    y = y_offset + row_idx * row_height
                    
                    # Рисуем границы ячейки
                    draw.rectangle([x, y, x + col_width, y + row_height], 
                                 outline='black', width=1)
                    
                    # Рисуем текст в ячейке
                    cell_text = cell.text.strip()[:30]  # Ограничиваем длину
                    if cell_text:
                        draw.text((x + 5, y + 5), cell_text, fill='black', font=font_normal)
            
            y_offset += len(table.rows) * row_height + 20
        
        # Сохраняем изображение
        image_filename = f"{base_name}_page_1.png"
        image_path = os.path.join(IMAGE_CACHE_DIR, image_filename)
        image.save(image_path, 'PNG', quality=95)
        
        print(f"✅ DOCX документ конвертирован в изображение: {image_path}")
        return image_path
        
    except Exception as e:
        print(f"❌ Ошибка конвертации DOCX: {e}")
        raise Exception(f"Не удалось конвертировать DOCX документ: {str(e)}")


def convert_xlsx_to_image(xlsx_path: str) -> str:
    """Конвертирует XLSX таблицу в изображение через openpyxl"""
    
    file_filename = os.path.basename(xlsx_path)
    base_name = os.path.splitext(file_filename)[0]
    
    print(f"📊 Конвертируем XLSX таблицу в изображение: {xlsx_path}")
    
    try:
        # Читаем XLSX файл
        workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
        worksheet = workbook.active  # Берем активный лист
        
        # Определяем размеры таблицы
        max_row = worksheet.max_row
        max_col = worksheet.max_column
        
        if max_row == 1 and max_col == 1:
            # Пустая таблица
            max_row = 10
            max_col = 5
        
        # Настройки изображения
        cell_width = 120
        cell_height = 30
        header_height = 35
        page_width = min(max_col * cell_width + 100, 1200)
        page_height = min(max_row * cell_height + 100, 1000)
        
        # Создаем изображение
        image = Image.new('RGB', (page_width, page_height), 'white')
        draw = ImageDraw.Draw(image)
        
        # Загружаем шрифты
        try:
            font_normal = ImageFont.truetype("arial.ttf", 10)
            font_bold = ImageFont.truetype("arialbd.ttf", 11)
        except:
            font_normal = ImageFont.load_default()
            font_bold = ImageFont.load_default()
        
        # Рисуем заголовок
        title = f"Excel таблица: {base_name}"
        draw.text((20, 20), title, fill='black', font=font_bold)
        
        y_start = 60
        
        # Рисуем таблицу
        for row_idx in range(1, min(max_row + 1, 25)):  # Ограничиваем количество строк
            for col_idx in range(1, min(max_col + 1, 10)):  # Ограничиваем количество колонок
                x = 20 + (col_idx - 1) * cell_width
                y = y_start + (row_idx - 1) * cell_height
                
                # Рисуем границы ячейки
                draw.rectangle([x, y, x + cell_width, y + cell_height], 
                             outline='black', width=1)
                
                # Получаем значение ячейки
                cell = worksheet.cell(row=row_idx, column=col_idx)
                cell_value = str(cell.value) if cell.value is not None else ""
                
                # Обрезаем длинный текст
                if len(cell_value) > 15:
                    cell_value = cell_value[:12] + "..."
                
                # Выбираем шрифт (жирный для первой строки)
                font = font_bold if row_idx == 1 else font_normal
                text_color = 'navy' if row_idx == 1 else 'black'
                
                # Рисуем текст в ячейке
                if cell_value:
                    draw.text((x + 5, y + 8), cell_value, fill=text_color, font=font)
        
        # Добавляем информацию о таблице
        info_text = f"Размер: {max_row} строк × {max_col} колонок"
        draw.text((20, page_height - 40), info_text, fill='gray', font=font_normal)
        
        # Сохраняем изображение
        image_filename = f"{base_name}_page_1.png"
        image_path = os.path.join(IMAGE_CACHE_DIR, image_filename)
        image.save(image_path, 'PNG', quality=95)
        
        print(f"✅ XLSX таблица конвертирована в изображение: {image_path}")
        return image_path
        
    except Exception as e:
        print(f"❌ Ошибка конвертации XLSX: {e}")
        raise Exception(f"Не удалось конвертировать XLSX таблицу: {str(e)}")


def convert_office_to_images(file_path: str) -> List[str]:
    """Конвертирует Office файлы (DOCX, XLSX) в изображения через Python библиотеки"""
    
    file_filename = os.path.basename(file_path)
    file_extension = os.path.splitext(file_filename)[1].lower()
    
    if not OFFICE_CONVERSION_AVAILABLE:
        raise Exception("Библиотеки для конвертации Office файлов недоступны")
    
    try:
        if file_extension in ['.docx', '.doc']:
            image_path = convert_docx_to_image(file_path)
            return [image_path]
        elif file_extension in ['.xlsx', '.xls']:
            image_path = convert_xlsx_to_image(file_path)
            return [image_path]
        else:
            raise Exception(f"Неподдерживаемый формат файла: {file_extension}")
            
    except Exception as e:
        print(f"❌ Ошибка конвертации Office файла: {e}")
        raise Exception(f"Не удалось конвертировать документ: {str(e)}")

# Удалены все сложные PDF функции - теперь используем простые изображения!

@router.post("/requests/{request_id}/fields/{field_name}/files/upload", response_model=List[RequestFileResponse])
async def upload_field_files(
    request_id: int,
    field_name: str,
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Загрузка файлов к полю формы заявки"""
    
    print(f"🔄 Backend: Получен запрос на загрузку файлов:")
    print(f"   Request ID: {request_id}")
    print(f"   Field name: {field_name}")
    print(f"   Files count: {len(files) if files else 0}")
    print(f"   Current user: {current_user.email if current_user else 'None'}")
    
    try:
        # Проверяем, существует ли заявка
        print(f"🔍 Backend: Ищем заявку с ID {request_id}")
        request = db.query(Request).filter(Request.id == request_id).first()
        if not request:
            print(f"❌ Backend: Заявка {request_id} не найдена")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Заявка не найдена"
            )
        
        print(f"✅ Backend: Заявка найдена: {request.title}, автор: {request.author_id}")
        
        # Проверяем права доступа (автор или ответственный)
        print(f"🔐 Backend: Проверяем права доступа")
        print(f"   Автор заявки: {request.author_id}")
        print(f"   Ответственный: {request.assignee_id}")
        print(f"   Текущий пользователь: {current_user.id}")
        
        if request.author_id != current_user.id and request.assignee_id != current_user.id:
            print(f"❌ Backend: Недостаточно прав для загрузки файлов")
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Недостаточно прав для загрузки файлов к этой заявке"
            )
        
        print(f"✅ Backend: Права доступа проверены успешно")
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Backend: Неожиданная ошибка при проверке заявки: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка сервера: {str(e)}"
        )
    
    uploaded_files = []
    
    print(f"📁 Backend: Начинаем обработку {len(files)} файлов")
    
    for i, file in enumerate(files, 1):
        print(f"📄 Backend: Обрабатываем файл {i}/{len(files)}: {file.filename}")
        print(f"   Размер: {file.size if hasattr(file, 'size') else 'неизвестно'}")
        print(f"   Тип: {file.content_type}")
        
        try:
            # Проверяем размер файла (максимум 10MB)
            content = await file.read()
            file_size = len(content)
            print(f"   Реальный размер: {file_size} байт")
            
            if file_size > 10 * 1024 * 1024:  # 10MB
                print(f"❌ Backend: Файл {file.filename} слишком большой: {file_size} байт")
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=f"Файл {file.filename} слишком большой. Максимальный размер: 10MB"
                )
            
            # Генерируем уникальное имя файла
            file_extension = os.path.splitext(file.filename)[1]
            unique_filename = f"{uuid.uuid4()}{file_extension}"
            file_path = os.path.join(UPLOAD_DIR, unique_filename)
            print(f"💾 Backend: Сохраняем как {unique_filename}")
            
            # Сохраняем файл
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(content)
            print(f"✅ Backend: Файл сохранен на диск")
            
            # Сохраняем информацию о файле в БД
            db_file = RequestFile(
                request_id=request_id,
                field_name=field_name,
                filename=file.filename,
                file_path=file_path,
                file_size=file_size,
                content_type=file.content_type or "application/octet-stream",
                uploaded_by=current_user.id
            )
            
            db.add(db_file)
            db.commit()
            db.refresh(db_file)
            print(f"✅ Backend: Информация о файле сохранена в БД с ID {db_file.id}")
            
            uploaded_files.append(db_file)
            
        except Exception as e:
            print(f"❌ Backend: Ошибка при обработке файла {file.filename}: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Ошибка при обработке файла {file.filename}: {str(e)}"
            )
    
    return uploaded_files

@router.get("/requests/{request_id}/fields/{field_name}/files", response_model=List[RequestFileResponse])
async def get_field_files(
    request_id: int,
    field_name: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Получение списка файлов поля"""
    
    # Проверяем, существует ли заявка
    request = db.query(Request).filter(Request.id == request_id).first()
    if not request:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Заявка не найдена"
        )
    
    # Проверяем права доступа (автор, ответственный или имеющий роль администратора/модератора)
    user_roles = current_user.roles
    has_admin_rights = any(role in ["admin", "moderator"] for role in user_roles)
    
    if (request.author_id != current_user.id and 
        request.assignee_id != current_user.id and 
        not has_admin_rights):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Недостаточно прав для просмотра файлов этой заявки"
        )
    
    files = db.query(RequestFile).filter(
        RequestFile.request_id == request_id,
        RequestFile.field_name == field_name
    ).all()
    return files

@router.get("/requests/{request_id}/files", response_model=List[RequestFileResponse])
async def get_all_request_files(
    request_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Получение всех файлов заявки (для отображения в деталях)"""
    
    # Проверяем, существует ли заявка
    request = db.query(Request).filter(Request.id == request_id).first()
    if not request:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Заявка не найдена"
        )
    
    # Проверяем права доступа (автор, ответственный или имеющий роль администратора/модератора)
    user_roles = current_user.roles
    has_admin_rights = any(role in ["admin", "moderator"] for role in user_roles)
    
    if (request.author_id != current_user.id and 
        request.assignee_id != current_user.id and 
        not has_admin_rights):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Недостаточно прав для просмотра файлов этой заявки"
        )
    
    files = db.query(RequestFile).filter(RequestFile.request_id == request_id).all()
    return files

@router.get("/files/{file_id}/download")
async def download_file(
    file_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Скачивание файла"""
    
    # Находим файл
    file_record = db.query(RequestFile).filter(RequestFile.id == file_id).first()
    if not file_record:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )
    
    # Проверяем права доступа к заявке
    request = file_record.request
    user_roles = current_user.roles
    has_admin_rights = any(role in ["admin", "moderator"] for role in user_roles)
    
    if (request.author_id != current_user.id and 
        request.assignee_id != current_user.id and 
        not has_admin_rights):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Недостаточно прав для скачивания этого файла"
        )
    
    # Проверяем, существует ли файл на диске
    if not os.path.exists(file_record.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден на сервере"
        )
    
    return FileResponse(
        path=file_record.file_path,
        filename=file_record.filename,
        media_type=file_record.content_type
    )

@router.get("/files/{file_id}/preview")
async def preview_file(
    file_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Предварительный просмотр файла"""
    
    # Находим файл
    file_record = db.query(RequestFile).filter(RequestFile.id == file_id).first()
    if not file_record:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )
    
    # Проверяем права доступа к заявке
    request = file_record.request
    user_roles = current_user.roles
    has_admin_rights = any(role in ["admin", "moderator"] for role in user_roles)
    
    if (request.author_id != current_user.id and 
        request.assignee_id != current_user.id and 
        not has_admin_rights):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Недостаточно прав для просмотра этого файла"
        )
    
    # Проверяем, существует ли файл на диске
    if not os.path.exists(file_record.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден на сервере"
        )
    
    return FileResponse(
        path=file_record.file_path,
        media_type=file_record.content_type
    )

@router.get("/files/{file_id}/preview-images")
async def preview_file_as_images(
    file_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Предварительный просмотр DOCX файла как изображения страниц - БЕЗ ВОДЯНЫХ ЗНАКОВ!"""
    
    # Находим файл
    file_record = db.query(RequestFile).filter(RequestFile.id == file_id).first()
    if not file_record:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )
    
    # Проверяем права доступа к заявке
    request = file_record.request
    user_roles = current_user.roles
    has_admin_rights = any(role in ["admin", "moderator"] for role in user_roles)
    
    if (request.author_id != current_user.id and 
        request.assignee_id != current_user.id and 
        not has_admin_rights):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Недостаточно прав для просмотра этого файла"
        )
    
    # Проверяем, существует ли файл на диске
    if not os.path.exists(file_record.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден на сервере"
        )
    
    # Проверяем, что это Office файл (DOCX или XLSX)
    supported_types = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',       # .xlsx
        'application/msword',                                                       # .doc
        'application/vnd.ms-excel'                                                  # .xls
    ]
    supported_extensions = ['.docx', '.doc', '.xlsx', '.xls']
    
    is_supported = (
        file_record.content_type in supported_types or
        'word' in file_record.content_type or
        'excel' in file_record.content_type or
        'spreadsheet' in file_record.content_type or
        any(file_record.filename.lower().endswith(ext) for ext in supported_extensions)
    )
    
    if not is_supported:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Конвертация в изображения поддерживается только для DOCX и XLSX файлов"
        )
    
    # Проверяем доступность библиотек для работы с изображениями
    if not OFFICE_CONVERSION_AVAILABLE:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Конвертация в изображения недоступна. Используйте стандартный просмотр."
        )
    
    try:
        # Конвертируем DOCX в изображения
        image_paths = convert_office_to_images(file_record.file_path)
        
        # Возвращаем первое изображение (первую страницу)
        # Frontend может запрашивать остальные страницы отдельно
        if image_paths and os.path.exists(image_paths[0]):
            return FileResponse(
                path=image_paths[0],
                media_type="image/png",
                filename=f"{os.path.splitext(file_record.filename)[0]}_page_1.png"
            )
        else:
            raise Exception("Не удалось создать изображения страниц")
        
    except Exception as e:
        print(f"❌ Ошибка конвертации файла {file_id} в изображения: {e}")
        
        # Возвращаем информативное сообщение об ошибке
        error_msg = "Конвертация в изображения временно недоступна"
        if "библиотеки" in str(e):
            error_msg = "Библиотеки для работы с изображениями недоступны"
        elif "docx2python" in str(e):
            error_msg = "Ошибка чтения DOCX файла"
        
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=f"{error_msg}. Используйте стандартный просмотр."
        )

@router.delete("/files/{file_id}")
async def delete_file(
    file_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Удаление файла"""
    
    # Находим файл
    file_record = db.query(RequestFile).filter(RequestFile.id == file_id).first()
    if not file_record:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )
    
    # Проверяем права доступа (только автор файла или администратор)
    user_roles = current_user.roles
    has_admin_rights = any(role in ["admin", "moderator"] for role in user_roles)
    
    if file_record.uploaded_by != current_user.id and not has_admin_rights:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Недостаточно прав для удаления этого файла"
        )
    
    # Удаляем файл с диска
    if os.path.exists(file_record.file_path):
        os.remove(file_record.file_path)
    
    # Удаляем запись из БД
    db.delete(file_record)
    db.commit()
    
    return {"message": "Файл успешно удален"} 